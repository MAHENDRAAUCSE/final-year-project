from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

BASE_DIR = Path(__file__).resolve().parents[1]
REPORT_MD = BASE_DIR / "reports" / "project_full_report.md"
PLOTS_DIR = BASE_DIR / "plots"
DATASET_CSV = BASE_DIR / "data" / "Rajahmundry_STP_Daily_Synthetic_2020_2023.csv"
OUTPUT_DOCX = BASE_DIR / "reports" / "Final_Project_Report.docx"
GENERATED_IMAGE = PLOTS_DIR / "06_monthly_avg_overview.png"

TEAM_NAMES = ["Mahendra", "Indra", "Shyam", "Bhanu", "Tanusri"]
ROLL_NUMBERS = [
    "322506402094",
    "322506402115",
    "322506402125",
    "322506402134",
    "322506402102",
]


def generate_extra_plot() -> None:
    """Generate a compact additional chart for the final document."""
    if GENERATED_IMAGE.exists():
        return

    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    df = pd.read_csv(DATASET_CSV)
    df["Date"] = pd.to_datetime(df["Date"])
    df["Month"] = df["Date"].dt.to_period("M").astype(str)

    source_columns = {
        "BOD": "BOD (mg/L)",
        "COD": "COD (mg/L)",
        "TSS": "TSS (mg/L)",
        "TN": "TN (mg/L)",
        "TP": "TP (mg/L)",
        "PH": "PH",
        "DO": "DO (mg/L)",
    }

    monthly = (
        df.groupby("Month")[list(source_columns.values())]
        .mean()
        .tail(18)
        .rename(columns={v: k for k, v in source_columns.items()})
    )

    plt.figure(figsize=(12, 6))
    for col in ["BOD", "COD", "DO", "PH"]:
        plt.plot(monthly.index, monthly[col], marker="o", linewidth=1.8, label=col)

    plt.xticks(rotation=45, ha="right")
    plt.title("Monthly Mean Trends (Last 18 Months)")
    plt.xlabel("Month")
    plt.ylabel("Average Value")
    plt.grid(True, alpha=0.25)
    plt.legend()
    plt.tight_layout()
    plt.savefig(GENERATED_IMAGE, dpi=220)
    plt.close()


def set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_cover_page(document: Document) -> None:
    p = document.add_paragraph("ANDHRA UNIVERSITY")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(20)

    p = document.add_paragraph("Department of Computer Science Engineering")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)

    p = document.add_paragraph("PROJECT REPORT")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(18)

    p = document.add_paragraph(
        "Smart STP Prediction System for Wastewater Quality Forecasting\n"
        "Using CNN-LSTM-Attention Model"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)

    document.add_paragraph("")
    document.add_paragraph("")

    p = document.add_paragraph("Submitted By")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Name"
    hdr[1].text = "Roll Number"

    for name, roll in zip(TEAM_NAMES, ROLL_NUMBERS):
        row = table.add_row().cells
        row[0].text = name
        row[1].text = roll

    document.add_paragraph("")

    for line in [
        "College: Andhra University",
        "University: Andhra University",
        "Branch: Computer Science Engineering",
        "Guide: Lavanya Kumari",
        "Year: Final Year",
        f"Date Generated: {datetime.now().strftime('%d-%m-%Y')}",
    ]:
        p = document.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if line.startswith(("College", "University", "Branch", "Guide", "Year")):
            p.runs[0].bold = True


def add_markdown_content(document: Document, lines: Iterable[str]) -> None:
    for raw_line in lines:
        line = raw_line.rstrip("\n")
        if not line.strip():
            document.add_paragraph("")
            continue

        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
            continue

        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
            continue

        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
            continue

        if line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue

        if line[:3].isdigit() and line[1:3] == ". ":
            document.add_paragraph(line[3:].strip(), style="List Number")
            continue

        document.add_paragraph(line)


def add_images_section(document: Document) -> None:
    document.add_page_break()
    document.add_heading("Project Visualizations", level=1)

    image_order = [
        "01_training_history.png",
        "02_predictions_vs_actual.png",
        "03_residuals.png",
        "04_metrics_summary.png",
        "05_actual_vs_predicted_scatter.png",
        "cod_02_predictions_vs_actual.png",
        "do_02_predictions_vs_actual.png",
        "ph_02_predictions_vs_actual.png",
        "06_monthly_avg_overview.png",
    ]

    for image_name in image_order:
        image_path = PLOTS_DIR / image_name
        if not image_path.exists():
            continue
        document.add_heading(image_name.replace("_", " ").replace(".png", "").title(), level=2)
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(6.4))


def main() -> None:
    if not REPORT_MD.exists():
        raise FileNotFoundError(f"Missing report source file: {REPORT_MD}")

    generate_extra_plot()

    document = Document()
    set_default_font(document)

    add_cover_page(document)
    document.add_section(start_type=WD_SECTION_START.NEW_PAGE)

    lines = REPORT_MD.read_text(encoding="utf-8").splitlines()
    add_markdown_content(document, lines)
    add_images_section(document)

    document.save(OUTPUT_DOCX)
    print(f"DOCX generated successfully: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
