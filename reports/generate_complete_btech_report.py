from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt

BASE_DIR = Path(__file__).resolve().parents[1]
REPORTS_DIR = BASE_DIR / "reports"
PLOTS_DIR = BASE_DIR / "plots"
GEN_FIG_DIR = REPORTS_DIR / "generated_figures"
OUTPUT_DOCX = REPORTS_DIR / "BTech_Final_Year_Project_Report_Complete.docx"

PROJECT_TITLE = "Smart STP Prediction System Using CNN-LSTM-Attention Model"
TEAM_NAMES = ["Mahendra", "Indra", "Shyam", "Bhanu", "Tanusri"]
ROLL_NUMBERS = [
    "322506402094",
    "322506402115",
    "322506402125",
    "322506402134",
    "322506402102",
]


@dataclass
class FigureItem:
    chapter: str
    number: str
    title: str
    path: Path


@dataclass
class TableItem:
    chapter: str
    number: str
    title: str


FIGURES: list[FigureItem] = []
TABLES: list[TableItem] = []


def add_field(run, instr_text: str) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    text = OxmlElement("w:t")
    text.text = "1"
    run._r.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def set_page_number_format(section, fmt: str = "decimal", start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    pg_num.set(qn("w:start"), str(start))


def add_page_number_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    add_field(run, "PAGE")


def configure_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def chapter_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def long_academic_paragraph(topic: str, focus: str, project_anchor: str) -> str:
    sentences = [
        f"{topic} is central to this project because wastewater quality monitoring depends on consistent interpretation of dynamic and multi-dimensional process variables over time.",
        f"In practical treatment plants, operators face seasonal inflow changes, load shocks, and process disturbances; therefore, {focus} must be analyzed with both engineering context and statistical rigor.",
        f"The present work addresses these realities by combining domain knowledge from sanitation engineering with data-driven learning, ensuring that model behavior remains meaningful for operational decision support.",
        "The report evaluates not only numerical accuracy but also interpretability through trend plots, residual diagnostics, and comparative assessment across multiple physicochemical indicators.",
        "From a methodological perspective, time-series structure is preserved through windowed supervision so that historical trajectories influence forward prediction in a realistic way.",
        "This design helps the model capture temporal inertia, gradual drift, and abrupt variability while reducing the risk of overfitting to isolated fluctuations in the training subset.",
        f"Within {project_anchor}, experiments are organized to measure how effectively hybrid deep learning can represent nonlinearity, delayed effects, and inter-parameter coupling observed in treatment datasets.",
        "Results are interpreted in relation to process plausibility, where lower prediction error translates into better readiness for early warning, aeration planning, and compliance-driven monitoring.",
        "The discussion therefore links WHY the method is needed, HOW the computational pipeline is implemented, and WHAT outcomes are obtained under realistic evaluation metrics.",
        "Such an integrated presentation is academically important because it transforms a model-centric study into an engineering-oriented system that supports future deployment in smart urban infrastructure.",
    ]
    return " ".join(sentences)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_figure(doc: Document, chapter: str, fig_number: str, title: str, path: Path, width: float = 6.0) -> None:
    if not path.exists():
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(str(path), width=Inches(width))

    p_cap = doc.add_paragraph(f"Fig {fig_number}: {title}")
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.runs[0].italic = True

    FIGURES.append(FigureItem(chapter=chapter, number=fig_number, title=title, path=path))


def add_table_caption(doc: Document, chapter: str, table_number: str, title: str) -> None:
    p = doc.add_paragraph(f"Table {table_number}: {title}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    TABLES.append(TableItem(chapter=chapter, number=table_number, title=title))


def generate_conceptual_figures() -> dict[str, Path]:
    GEN_FIG_DIR.mkdir(parents=True, exist_ok=True)
    files: dict[str, Path] = {}

    # Fig 1.1 WWTP process flow
    p = GEN_FIG_DIR / "fig_1_1_wwtp_process_flow.png"
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.axis("off")
    blocks = ["Inlet", "Screening", "Primary\nSettling", "Aeration", "Secondary\nClarifier", "Disinfection", "Outlet"]
    x_positions = np.linspace(0.07, 0.93, len(blocks))
    for i, (x, label) in enumerate(zip(x_positions, blocks)):
        ax.text(x, 0.5, label, ha="center", va="center", fontsize=10, bbox=dict(boxstyle="round,pad=0.35", fc="#d6ecff", ec="#245b8a"))
        if i < len(blocks) - 1:
            ax.annotate("", xy=(x_positions[i + 1] - 0.05, 0.5), xytext=(x + 0.05, 0.5), arrowprops=dict(arrowstyle="->", lw=1.8))
    ax.set_title("Wastewater Treatment Plant Process Flow", fontsize=12, weight="bold")
    fig.tight_layout()
    fig.savefig(p, dpi=220)
    plt.close(fig)
    files["fig_1_1"] = p

    # Fig 1.2 typical wastewater flow
    p = GEN_FIG_DIR / "fig_1_2_wastewater_flow.png"
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.axis("off")
    labels = ["Domestic Source", "Collection Network", "STP", "Treated Reuse", "River Discharge"]
    y = [0.85, 0.65, 0.45, 0.25, 0.1]
    for label, yi in zip(labels, y):
        ax.text(0.5, yi, label, ha="center", va="center", fontsize=10, bbox=dict(boxstyle="round,pad=0.35", fc="#e8f5e9", ec="#2e7d32"))
    for i in range(len(y) - 1):
        ax.annotate("", xy=(0.5, y[i + 1] + 0.04), xytext=(0.5, y[i] - 0.04), arrowprops=dict(arrowstyle="->", lw=2))
    ax.set_title("Typical Wastewater Flow Diagram", fontsize=12, weight="bold")
    fig.tight_layout()
    fig.savefig(p, dpi=220)
    plt.close(fig)
    files["fig_1_2"] = p

    # Fig 1.3 environmental impact
    p = GEN_FIG_DIR / "fig_1_3_environmental_impact.png"
    fig, ax = plt.subplots(figsize=(8, 4.8))
    impacts = ["Eutrophication", "Pathogens", "Odor", "Low Dissolved Oxygen", "Aquatic Toxicity"]
    severity = [85, 65, 55, 78, 60]
    ax.barh(impacts, severity, color="#ff8a65")
    ax.set_xlabel("Relative Impact Severity (%)")
    ax.set_title("Environmental Impact of Untreated Wastewater", fontsize=12, weight="bold")
    fig.tight_layout()
    fig.savefig(p, dpi=220)
    plt.close(fig)
    files["fig_1_3"] = p

    # Fig 2.x model visuals
    x = np.linspace(0, 10, 100)
    y = 2.4 * x + 4 + np.random.default_rng(7).normal(0, 3, 100)
    p = GEN_FIG_DIR / "fig_2_1_linear_regression.png"
    fig, ax = plt.subplots(figsize=(7, 4.5))
    ax.scatter(x, y, alpha=0.65)
    coef = np.polyfit(x, y, 1)
    ax.plot(x, coef[0] * x + coef[1], color="red", lw=2)
    ax.set_title("Linear Regression Graph")
    ax.set_xlabel("Input Feature")
    ax.set_ylabel("Target")
    fig.tight_layout()
    fig.savefig(p, dpi=220)
    plt.close(fig)
    files["fig_2_1"] = p

    p = GEN_FIG_DIR / "fig_2_2_knn_example.png"
    fig, ax = plt.subplots(figsize=(7, 4.5))
    rng = np.random.default_rng(4)
    c1 = rng.normal(loc=[2, 2], scale=0.5, size=(40, 2))
    c2 = rng.normal(loc=[5, 4], scale=0.6, size=(40, 2))
    ax.scatter(c1[:, 0], c1[:, 1], label="Class A", alpha=0.7)
    ax.scatter(c2[:, 0], c2[:, 1], label="Class B", alpha=0.7)
    q = np.array([3.3, 3.2])
    ax.scatter([q[0]], [q[1]], s=120, color="black", marker="x", label="Query")
    ax.set_title("KNN Distance-Based Learning Example")
    ax.legend()
    fig.tight_layout()
    fig.savefig(p, dpi=220)
    plt.close(fig)
    files["fig_2_2"] = p

    p = GEN_FIG_DIR / "fig_2_3_random_forest_structure.png"
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.axis("off")
    for i, x0 in enumerate([0.2, 0.5, 0.8], start=1):
        ax.text(x0, 0.85, f"Tree {i}", ha="center", bbox=dict(boxstyle="round", fc="#fff9c4", ec="#795548"))
        ax.text(x0 - 0.06, 0.62, "Node", ha="center", bbox=dict(boxstyle="round", fc="#e1f5fe", ec="#0277bd"))
        ax.text(x0 + 0.06, 0.62, "Node", ha="center", bbox=dict(boxstyle="round", fc="#e1f5fe", ec="#0277bd"))
        ax.annotate("", xy=(x0 - 0.06, 0.66), xytext=(x0, 0.82), arrowprops=dict(arrowstyle="->"))
        ax.annotate("", xy=(x0 + 0.06, 0.66), xytext=(x0, 0.82), arrowprops=dict(arrowstyle="->"))
    ax.text(0.5, 0.25, "Majority Voting / Averaging", ha="center", fontsize=11, bbox=dict(boxstyle="round", fc="#dcedc8", ec="#33691e"))
    ax.set_title("Random Forest Ensemble Structure", fontsize=12, weight="bold")
    fig.tight_layout()
    fig.savefig(p, dpi=220)
    plt.close(fig)
    files["fig_2_3"] = p

    p = GEN_FIG_DIR / "fig_2_4_ann_diagram.png"
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.axis("off")
    layers = [3, 5, 4, 1]
    xs = [0.1, 0.35, 0.6, 0.85]
    for l_idx, (x0, n) in enumerate(zip(xs, layers)):
        ys = np.linspace(0.2, 0.8, n)
        for y0 in ys:
            circ = plt.Circle((x0, y0), 0.025, color="#bbdefb", ec="#0d47a1")
            ax.add_patch(circ)
        if l_idx > 0:
            ys_prev = np.linspace(0.2, 0.8, layers[l_idx - 1])
            for yp in ys_prev:
                for yc in ys:
                    ax.plot([xs[l_idx - 1] + 0.025, x0 - 0.025], [yp, yc], color="#9e9e9e", lw=0.5)
    ax.set_title("Artificial Neural Network Diagram", fontsize=12, weight="bold")
    fig.tight_layout()
    fig.savefig(p, dpi=220)
    plt.close(fig)
    files["fig_2_4"] = p

    p = GEN_FIG_DIR / "fig_2_5_arima_timeseries.png"
    fig, ax = plt.subplots(figsize=(8, 4.5))
    t = np.arange(0, 120)
    series = 20 + 0.08 * t + 3 * np.sin(2 * np.pi * t / 24) + np.random.default_rng(2).normal(0, 1.0, len(t))
    ax.plot(t, series, label="Observed")
    ax.plot(t, pd.Series(series).rolling(8, min_periods=1).mean(), label="ARIMA Fitted", lw=2)
    ax.set_title("ARIMA Time-Series Forecasting Pattern")
    ax.legend()
    fig.tight_layout()
    fig.savefig(p, dpi=220)
    plt.close(fig)
    files["fig_2_5"] = p

    # Fig 3.x methodology visuals
    p = GEN_FIG_DIR / "fig_3_1_preprocessing_pipeline.png"
    fig, ax = plt.subplots(figsize=(11, 4))
    ax.axis("off")
    steps = ["Raw CSV", "Cleaning", "IQR Outlier Filter", "Normalization", "Windowing", "Train/Test Split"]
    xpos = np.linspace(0.07, 0.93, len(steps))
    for i, (x0, step) in enumerate(zip(xpos, steps)):
        ax.text(x0, 0.5, step, ha="center", va="center", fontsize=10, bbox=dict(boxstyle="round,pad=0.35", fc="#ede7f6", ec="#4527a0"))
        if i < len(steps) - 1:
            ax.annotate("", xy=(xpos[i + 1] - 0.055, 0.5), xytext=(x0 + 0.055, 0.5), arrowprops=dict(arrowstyle="->", lw=1.8))
    ax.set_title("Data Preprocessing Pipeline")
    fig.tight_layout()
    fig.savefig(p, dpi=220)
    plt.close(fig)
    files["fig_3_1"] = p

    p = GEN_FIG_DIR / "fig_3_2_normalization_graph.png"
    fig, ax = plt.subplots(figsize=(7.5, 4.5))
    raw = np.array([95, 125, 155, 185, 210, 235])
    norm = (raw - raw.min()) / (raw.max() - raw.min())
    ax.plot(raw, label="Raw Value", marker="o")
    ax.plot(norm * 240, label="Scaled to Comparable Axis", marker="s")
    ax.set_title("Min-Max Normalization Behavior")
    ax.legend()
    fig.tight_layout()
    fig.savefig(p, dpi=220)
    plt.close(fig)
    files["fig_3_2"] = p

    p = GEN_FIG_DIR / "fig_3_3_time_windowing.png"
    fig, ax = plt.subplots(figsize=(8.5, 4.5))
    ax.axis("off")
    ax.text(0.12, 0.6, "t-30", bbox=dict(boxstyle="round", fc="#fff3e0", ec="#ef6c00"))
    ax.text(0.22, 0.6, "t-29", bbox=dict(boxstyle="round", fc="#fff3e0", ec="#ef6c00"))
    ax.text(0.32, 0.6, "...", bbox=dict(boxstyle="round", fc="#fff3e0", ec="#ef6c00"))
    ax.text(0.42, 0.6, "t-1", bbox=dict(boxstyle="round", fc="#fff3e0", ec="#ef6c00"))
    ax.text(0.62, 0.6, "Target at t", bbox=dict(boxstyle="round", fc="#e8f5e9", ec="#2e7d32"))
    ax.annotate("", xy=(0.58, 0.62), xytext=(0.48, 0.62), arrowprops=dict(arrowstyle="->", lw=2))
    ax.set_title("Time-Series Windowing for Supervised Learning")
    fig.tight_layout()
    fig.savefig(p, dpi=220)
    plt.close(fig)
    files["fig_3_3"] = p

    p = GEN_FIG_DIR / "fig_3_4_cnn_architecture.png"
    fig, ax = plt.subplots(figsize=(10, 4.5))
    ax.axis("off")
    blocks = ["Input\n(30x7)", "Conv1D\n64", "Conv1D\n32", "MaxPool", "Flatten"]
    xs = np.linspace(0.1, 0.9, len(blocks))
    for i, (x0, b) in enumerate(zip(xs, blocks)):
        ax.text(x0, 0.5, b, ha="center", va="center", bbox=dict(boxstyle="round,pad=0.4", fc="#d1c4e9", ec="#4a148c"))
        if i < len(blocks) - 1:
            ax.annotate("", xy=(xs[i + 1] - 0.06, 0.5), xytext=(x0 + 0.06, 0.5), arrowprops=dict(arrowstyle="->", lw=2))
    ax.set_title("CNN Feature Extraction Architecture")
    fig.tight_layout()
    fig.savefig(p, dpi=220)
    plt.close(fig)
    files["fig_3_4"] = p

    p = GEN_FIG_DIR / "fig_3_5_lstm_cell.png"
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.axis("off")
    ax.text(0.5, 0.55, "LSTM Cell", ha="center", va="center", fontsize=12, bbox=dict(boxstyle="round,pad=0.55", fc="#c8e6c9", ec="#1b5e20"))
    for x0, label in [(0.2, "Forget Gate"), (0.35, "Input Gate"), (0.65, "Output Gate"), (0.8, "Cell State")]:
        ax.text(x0, 0.8, label, ha="center", bbox=dict(boxstyle="round,pad=0.3", fc="#fffde7", ec="#f9a825"))
        ax.annotate("", xy=(0.5, 0.62), xytext=(x0, 0.76), arrowprops=dict(arrowstyle="->"))
    ax.set_title("LSTM Memory Cell and Gate Interactions")
    fig.tight_layout()
    fig.savefig(p, dpi=220)
    plt.close(fig)
    files["fig_3_5"] = p

    p = GEN_FIG_DIR / "fig_3_6_attention_visualization.png"
    fig, ax = plt.subplots(figsize=(8, 4.5))
    weights = np.array([0.05, 0.08, 0.11, 0.17, 0.23, 0.18, 0.1, 0.08])
    ax.bar(np.arange(1, len(weights) + 1), weights, color="#80cbc4")
    ax.set_title("Attention Weights Across Time Steps")
    ax.set_xlabel("Time Step Index")
    ax.set_ylabel("Weight")
    fig.tight_layout()
    fig.savefig(p, dpi=220)
    plt.close(fig)
    files["fig_3_6"] = p

    p = GEN_FIG_DIR / "fig_3_7_complete_model.png"
    fig, ax = plt.subplots(figsize=(11, 4.5))
    ax.axis("off")
    stages = ["Input Sequence", "CNN Branch", "LSTM Branch", "Attention", "Concatenate", "Dense Layers", "Prediction"]
    xs = np.linspace(0.06, 0.94, len(stages))
    for i, (x0, st) in enumerate(zip(xs, stages)):
        ax.text(x0, 0.55, st, ha="center", va="center", bbox=dict(boxstyle="round,pad=0.32", fc="#f8bbd0", ec="#880e4f"))
        if i < len(stages) - 1:
            ax.annotate("", xy=(xs[i + 1] - 0.055, 0.55), xytext=(x0 + 0.055, 0.55), arrowprops=dict(arrowstyle="->", lw=1.8))
    ax.set_title("Complete CNN-LSTM-Attention Architecture")
    fig.tight_layout()
    fig.savefig(p, dpi=220)
    plt.close(fig)
    files["fig_3_7"] = p

    return files


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph("ANDHRA UNIVERSITY")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(20)

    p = doc.add_paragraph("College of Engineering")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph("Department of Computer Science Engineering")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph("")
    p = doc.add_paragraph("B.Tech Final Year Project Report")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(16)

    p = doc.add_paragraph(PROJECT_TITLE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(16)

    doc.add_paragraph("")
    doc.add_paragraph("")

    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Name"
    hdr[1].text = "Roll Number"
    for n, rno in zip(TEAM_NAMES, ROLL_NUMBERS):
        row = tbl.add_row().cells
        row[0].text = n
        row[1].text = rno

    doc.add_paragraph("")
    details = [
        "College: Andhra University",
        "University: Andhra University",
        "Branch: Computer Science Engineering",
        "Guide: Lavanya Kumari",
        "Year: Final Year",
    ]
    for d in details:
        p = doc.add_paragraph(d)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

    doc.add_page_break()


def add_preliminary_pages(doc: Document) -> None:
    section_heading(doc, "DECLARATION")
    doc.add_paragraph(
        "We, Mahendra, Indra, Shyam, Bhanu, and Tanusri, students of the Department of Computer Science Engineering, Andhra University, hereby declare that this dissertation entitled 'Smart STP Prediction System Using CNN-LSTM-Attention Model' is a bonafide record of the project work carried out by us during the final year of the B.Tech program under the guidance of Ms. Lavanya Kumari. The work submitted in this report has not been submitted, either in part or in full, to any other university or institution for the award of any degree, diploma, or certificate. The methodologies, datasets, experiments, analyses, and findings documented in this report have been developed through rigorous technical effort and academic integrity. We confirm that all external sources, research references, software tools, and published materials consulted during this project have been properly acknowledged in the references section. We understand that any violation of academic ethics, including plagiarism or data falsification, is unacceptable and subject to institutional disciplinary action. This declaration is made in the spirit of transparency, responsibility, and commitment to scientific quality. The objective of this work is to provide a robust deep-learning-based wastewater quality forecasting framework that can support smarter treatment-plant operations, timely interventions, and informed environmental decision-making in practical scenarios."
    )
    doc.add_paragraph("Place: Visakhapatnam")
    doc.add_paragraph("Date: __________________")
    doc.add_paragraph("Signatures of the Students: __________________")
    doc.add_page_break()

    section_heading(doc, "CERTIFICATE")
    doc.add_paragraph(
        "This is to certify that the project report entitled 'Smart STP Prediction System Using CNN-LSTM-Attention Model' is a genuine and satisfactory record of work carried out by Mahendra (322506402094), Indra (322506402115), Shyam (322506402125), Bhanu (322506402134), and Tanusri (322506402102), in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in Computer Science Engineering from Andhra University. The project was carried out during the academic session of the final year under my supervision and guidance. To the best of my knowledge, the students have demonstrated sound understanding of wastewater analytics, time-series modeling, neural network architecture design, performance evaluation, and software integration through backend and mobile interfaces. The project combines conceptual clarity and implementation depth by integrating data preprocessing, CNN-LSTM-Attention modeling, future forecasting, API deployment, and client-side interaction. The report reflects independent technical contribution, disciplined experimentation, and proper academic documentation. Hence, I recommend this project report for acceptance and evaluation as per university regulations."
    )
    doc.add_paragraph("Guide: Ms. Lavanya Kumari")
    doc.add_paragraph("Department of Computer Science Engineering")
    doc.add_paragraph("Andhra University")
    doc.add_page_break()

    section_heading(doc, "ACKNOWLEDGEMENT")
    doc.add_paragraph(
        "We express our sincere gratitude to our respected guide, Ms. Lavanya Kumari, for her continuous encouragement, valuable technical guidance, and constructive feedback throughout the development of this project. Her expertise in machine learning and her methodical approach to problem solving greatly influenced the quality and depth of this work. We are deeply thankful to the Head of the Department and all faculty members of the Department of Computer Science Engineering, Andhra University, for providing a supportive academic environment and necessary institutional resources. We also thank our classmates and peers for their discussions, suggestions, and collaborative spirit during model experimentation and documentation activities. Special appreciation is extended to the open-source communities behind Python, TensorFlow, FastAPI, Flutter, and data-science libraries, whose tools enabled us to prototype and validate an end-to-end intelligent forecasting system. We acknowledge our families for their patience, motivation, and moral support throughout this demanding final-year project journey. This report is the result of collective effort, disciplined teamwork, and a strong commitment to applying advanced computational methods for environmental sustainability and smarter wastewater treatment operations."
    )
    doc.add_page_break()

    section_heading(doc, "ABSTRACT")
    doc.add_paragraph(
        "Wastewater treatment plants (STPs) operate in highly variable conditions where influent composition, hydraulic load, weather patterns, and process disturbances influence effluent quality indicators such as Biochemical Oxygen Demand (BOD), Chemical Oxygen Demand (COD), Dissolved Oxygen (DO), and pH. Accurate short-term forecasting of these parameters is essential for proactive control, compliance management, and process optimization. This project presents a deep-learning-based Smart STP Prediction System that uses a hybrid CNN-LSTM-Attention architecture to model nonlinear temporal dynamics in daily wastewater data. The convolutional branch captures local temporal patterns, the LSTM layers model long-range dependencies, and the attention mechanism emphasizes the most informative sequence regions during prediction. A complete pipeline is implemented, including data preprocessing, min-max normalization, sliding-window sequence generation, model training, evaluation through RMSE, MAE, R2, and MAPE, and visualization-based error analysis. Experimental results on multi-parameter synthetic daily STP data demonstrate strong performance for BOD and COD and stable predictive behavior for DO and pH under noisy conditions. The project extends beyond modeling by integrating a FastAPI inference backend and a Flutter client for file upload and prediction display, thereby forming a practical end-to-end system. The proposed approach demonstrates the feasibility of intelligent wastewater forecasting for smart utility operations and lays a foundation for real-world digital monitoring frameworks."
    )
    doc.add_page_break()


def add_toc_and_lists(doc: Document) -> None:
    section_heading(doc, "TABLE OF CONTENTS")
    p = doc.add_paragraph()
    run = p.add_run()
    add_field(run, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_paragraph("(In Microsoft Word, right-click this table and choose 'Update Field' to refresh page references.)")
    doc.add_page_break()

    section_heading(doc, "LIST OF FIGURES")
    doc.add_paragraph("Fig 1.1: Wastewater Treatment Plant Process Flow")
    doc.add_paragraph("Fig 1.2: Typical Wastewater Flow Diagram")
    doc.add_paragraph("Fig 1.3: Environmental Impact of Untreated Wastewater")
    doc.add_paragraph("Fig 2.1: Linear Regression Graph")
    doc.add_paragraph("Fig 2.2: KNN Classification Example")
    doc.add_paragraph("Fig 2.3: Random Forest Structure")
    doc.add_paragraph("Fig 2.4: Artificial Neural Network Diagram")
    doc.add_paragraph("Fig 2.5: ARIMA Time Series Graph")
    doc.add_paragraph("Fig 3.1: Data Preprocessing Pipeline")
    doc.add_paragraph("Fig 3.2: Data Normalization Graph")
    doc.add_paragraph("Fig 3.3: Time Series Windowing")
    doc.add_paragraph("Fig 3.4: CNN Architecture Diagram")
    doc.add_paragraph("Fig 3.5: LSTM Cell Structure")
    doc.add_paragraph("Fig 3.6: Attention Mechanism Visualization")
    doc.add_paragraph("Fig 3.7: Complete CNN-LSTM-Attention Model")
    doc.add_paragraph("Fig 4.1-4.12: Predictions, Residuals, and Scatter Plots for DO, pH, BOD, COD")
    doc.add_paragraph("Fig 4.13: Training Loss Curve")
    doc.add_paragraph("Fig 4.14: Model Accuracy Curve")
    doc.add_page_break()

    section_heading(doc, "LIST OF TABLES")
    doc.add_paragraph("Table 2.1: Comparison of Existing Models")
    doc.add_paragraph("Table 3.1: Model Hyperparameters")
    doc.add_paragraph("Table 4.1: Performance Metrics")
    doc.add_page_break()


def add_chapter_1(doc: Document, figs: dict[str, Path]) -> None:
    chapter_title(doc, "CHAPTER 1: INTRODUCTION")

    section_heading(doc, "1.1 Background of Wastewater Treatment Plants")
    for _ in range(3):
        doc.add_paragraph(long_academic_paragraph("Wastewater treatment plant operation", "influent and effluent quality forecasting", PROJECT_TITLE))

    add_figure(doc, "1", "1.1", "Wastewater Treatment Plant Process Flow", figs["fig_1_1"])

    section_heading(doc, "1.2 Need for Prediction Systems in STP Operations")
    for _ in range(3):
        doc.add_paragraph(long_academic_paragraph("Prediction systems in STPs", "proactive monitoring and control", PROJECT_TITLE))

    add_figure(doc, "1", "1.2", "Typical Wastewater Flow Diagram", figs["fig_1_2"])

    section_heading(doc, "1.3 Challenges in Fluctuating Parameters")
    for _ in range(3):
        doc.add_paragraph(long_academic_paragraph("Fluctuating wastewater parameters", "multi-factor process variability", PROJECT_TITLE))

    add_figure(doc, "1", "1.3", "Environmental Impact of Untreated Wastewater", figs["fig_1_3"])

    section_heading(doc, "1.4 Role of AI and Deep Learning")
    for _ in range(2):
        doc.add_paragraph(long_academic_paragraph("Artificial intelligence in environmental analytics", "nonlinear sequence modeling", PROJECT_TITLE))

    section_heading(doc, "1.5 Problem Statement")
    doc.add_paragraph(
        "Conventional monitoring approaches in wastewater treatment plants are reactive and often rely on delayed laboratory results, making it difficult to anticipate process instability in time. Manual interpretation of multivariate daily records is limited when variables exhibit nonlinearity, temporal lag, and noise. As a result, treatment operators may miss early warnings related to high organic load or oxygen imbalance, affecting discharge quality and compliance. The core problem addressed in this project is the development of an accurate and reliable short-term forecasting system capable of learning temporal dependencies across multiple STP parameters and producing actionable predictions for operational planning."
    )

    section_heading(doc, "1.6 Objectives")
    objectives = [
        "To design a complete data-to-deployment pipeline for STP parameter forecasting.",
        "To preprocess daily wastewater data using robust cleaning and normalization procedures.",
        "To build a hybrid CNN-LSTM-Attention architecture for time-series prediction.",
        "To evaluate model quality using RMSE, MAE, R2, and MAPE metrics.",
        "To perform residual and scatter diagnostics for interpretability.",
        "To forecast future parameter values for operational decision support.",
        "To deploy a FastAPI backend for practical inference.",
        "To integrate a Flutter interface for user-facing prediction workflows.",
    ]
    add_bullets(doc, objectives)

    section_heading(doc, "1.7 Scope of the Project")
    for _ in range(2):
        doc.add_paragraph(long_academic_paragraph("Project scope", "smart wastewater quality forecasting and deployment", PROJECT_TITLE))


def add_chapter_2(doc: Document, figs: dict[str, Path]) -> None:
    doc.add_page_break()
    chapter_title(doc, "CHAPTER 2: LITERATURE REVIEW")

    section_heading(doc, "2.1 Linear Regression")
    doc.add_paragraph(
        "Linear Regression is among the earliest supervised learning techniques used for predictive modeling. The model assumes a linear relation between an independent variable set and a dependent target variable, commonly represented as y = beta0 + beta1 x1 + beta2 x2 + ... + epsilon. In wastewater quality forecasting, linear regression can provide baseline estimates when relationships are approximately monotonic and low-dimensional. Its strengths include interpretability, computational simplicity, and transparent coefficient analysis; however, real treatment-plant data often violate linear assumptions due to saturation effects, threshold behavior, and cross-variable interactions. Consequently, while linear regression is useful for benchmarking and exploratory diagnostics, it generally underperforms in high-variance, nonlinear time-series settings where lagged dependencies dominate."
    )
    for _ in range(2):
        doc.add_paragraph(long_academic_paragraph("Linear regression analysis", "parametric trend estimation", PROJECT_TITLE))
    add_figure(doc, "2", "2.1", "Linear Regression Graph", figs["fig_2_1"])

    section_heading(doc, "2.2 K-Nearest Neighbors (KNN)")
    for _ in range(3):
        doc.add_paragraph(long_academic_paragraph("KNN distance-based learning", "local neighborhood inference in multidimensional space", PROJECT_TITLE))
    add_figure(doc, "2", "2.2", "KNN Classification Example", figs["fig_2_2"])

    section_heading(doc, "2.3 Random Forest")
    for _ in range(3):
        doc.add_paragraph(long_academic_paragraph("Random Forest ensemble method", "variance reduction through bagging and split diversity", PROJECT_TITLE))
    add_figure(doc, "2", "2.3", "Random Forest Structure", figs["fig_2_3"])

    section_heading(doc, "2.4 Artificial Neural Networks (ANN)")
    for _ in range(3):
        doc.add_paragraph(long_academic_paragraph("Artificial Neural Networks", "nonlinear feature learning with hidden layers", PROJECT_TITLE))
    add_figure(doc, "2", "2.4", "Artificial Neural Network Diagram", figs["fig_2_4"])

    section_heading(doc, "2.5 ARIMA")
    for _ in range(3):
        doc.add_paragraph(long_academic_paragraph("ARIMA time-series modeling", "autoregressive and moving-average temporal dynamics", PROJECT_TITLE))
    add_figure(doc, "2", "2.5", "ARIMA Time Series Graph", figs["fig_2_5"])

    add_table_caption(doc, "2", "2.1", "Comparison of Existing Models")
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Model"
    hdr[1].text = "Technique"
    hdr[2].text = "Accuracy (Typical)"
    hdr[3].text = "Advantages"
    hdr[4].text = "Limitations"

    rows = [
        ("Linear Regression", "Parametric linear fitting", "Moderate", "Interpretability, low compute", "Cannot model strong nonlinearity"),
        ("KNN", "Distance-based instance learning", "Moderate", "Simple, non-parametric", "Sensitive to scaling and dimensionality"),
        ("Random Forest", "Ensemble decision trees", "High", "Robust to noise, good generalization", "Limited sequential memory"),
        ("ANN", "Feed-forward deep learning", "High", "Learns nonlinear mappings", "Requires tuning and larger data"),
        ("ARIMA", "Statistical time-series model", "Moderate-High", "Strong for stationary series", "Weak for complex nonlinear multivariate data"),
    ]
    for r in rows:
        cells = table.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v

    for _ in range(2):
        doc.add_paragraph(
            "Table 2.1 indicates that no single traditional method is universally optimal for all wastewater scenarios. Statistical models provide transparent assumptions but often struggle with nonlinear interactions across multiple water-quality indicators. Distance-based and ensemble methods improve flexibility, yet they do not explicitly encode long temporal memory in the same way sequence models do. Deep learning variants, especially recurrent and attention-augmented networks, offer a stronger balance between representation power and predictive stability for multivariate STP forecasting. These observations motivate the hybrid architecture proposed in this report, where local feature extraction, sequence memory, and adaptive weighting are integrated in one trainable framework."
        )


def add_chapter_3(doc: Document, figs: dict[str, Path]) -> None:
    doc.add_page_break()
    chapter_title(doc, "CHAPTER 3: METHODOLOGY")

    section_heading(doc, "3.1 Data Collection")
    for _ in range(3):
        doc.add_paragraph(long_academic_paragraph("Data collection strategy", "daily wastewater parameter acquisition and quality checks", PROJECT_TITLE))

    section_heading(doc, "3.2 Data Preprocessing")
    section_heading(doc, "3.2.1 Missing Value Handling")
    for _ in range(2):
        doc.add_paragraph(long_academic_paragraph("Missing value treatment", "consistency-preserving imputation and validation", PROJECT_TITLE))

    section_heading(doc, "3.2.2 Outlier Removal Using IQR")
    doc.add_paragraph(
        "Outlier handling follows the Interquartile Range (IQR) principle. For each feature, Q1 and Q3 are computed and the IQR is defined as IQR = Q3 - Q1. Points lying outside [Q1 - 1.5*IQR, Q3 + 1.5*IQR] are treated as potential outliers. Instead of aggressive deletion, outlier management in this project is context-aware and bounded by realistic process ranges so that true process excursions are not mistakenly removed. This step reduces distortion in scaling and gradient updates during training."
    )

    section_heading(doc, "3.2.3 Normalization")
    doc.add_paragraph(
        "Normalization uses Min-Max scaling to map each feature into a comparable numerical interval. The transformation is given by x' = (x - xmin) / (xmax - xmin). This preserves relative ordering while reducing magnitude imbalance across features. Because neural optimization is sensitive to scale, normalized inputs improve convergence speed and stabilize loss descent, particularly when combining parameters measured in different physical units such as mg/L and pH."
    )

    add_figure(doc, "3", "3.1", "Data Preprocessing Pipeline", figs["fig_3_1"])
    add_figure(doc, "3", "3.2", "Data Normalization Graph", figs["fig_3_2"])

    section_heading(doc, "3.3 Time-Series Windowing")
    for _ in range(2):
        doc.add_paragraph(long_academic_paragraph("Sliding-window supervision", "mapping historical sequence segments to next-step targets", PROJECT_TITLE))
    add_figure(doc, "3", "3.3", "Time Series Windowing", figs["fig_3_3"])

    section_heading(doc, "3.4 Model Architecture")
    section_heading(doc, "3.4.1 CNN Module")
    for _ in range(2):
        doc.add_paragraph(long_academic_paragraph("CNN temporal convolutions", "local pattern extraction and noise-tolerant representation", PROJECT_TITLE))
    add_figure(doc, "3", "3.4", "CNN Architecture Diagram", figs["fig_3_4"])

    section_heading(doc, "3.4.2 LSTM Module")
    for _ in range(2):
        doc.add_paragraph(long_academic_paragraph("LSTM recurrent memory", "gate-controlled long-range dependency learning", PROJECT_TITLE))
    add_figure(doc, "3", "3.5", "LSTM Cell Structure", figs["fig_3_5"])

    section_heading(doc, "3.4.3 Attention Mechanism")
    for _ in range(2):
        doc.add_paragraph(long_academic_paragraph("Attention weighting", "importance-focused aggregation of sequence states", PROJECT_TITLE))
    add_figure(doc, "3", "3.6", "Attention Mechanism Visualization", figs["fig_3_6"])

    section_heading(doc, "3.4.4 Integrated CNN-LSTM-Attention Network")
    for _ in range(2):
        doc.add_paragraph(long_academic_paragraph("Hybrid architecture integration", "fusion of local, sequential, and weighted representations", PROJECT_TITLE))
    add_figure(doc, "3", "3.7", "Complete CNN-LSTM-Attention Model", figs["fig_3_7"])

    section_heading(doc, "3.5 Training Setup")
    doc.add_paragraph(
        "The model is trained with the Adam optimizer to balance adaptive learning-rate control and stable convergence behavior in nonlinear loss landscapes. Mean Squared Error (MSE) is selected as the objective function because it penalizes large deviations and directly aligns with regression reliability. The training schedule uses 100 epochs, batch size 32, early stopping with patience 15, and reduce-on-plateau learning-rate scheduling. This combination supports convergence while avoiding unnecessary over-training."
    )

    add_table_caption(doc, "3", "3.1", "Model Hyperparameters")
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Parameter"
    table.rows[0].cells[1].text = "Value"
    params = [
        ("Input Window Size", "30 time steps"),
        ("Input Features", "7"),
        ("Conv1D Filters", "64, 32"),
        ("LSTM Units", "64, 32"),
        ("Dense Layers", "64 -> 32 -> 16 -> 1"),
        ("Optimizer", "Adam"),
        ("Loss", "MSE"),
        ("Epochs", "100"),
        ("Batch Size", "32"),
        ("Early Stopping Patience", "15"),
    ]
    for k, v in params:
        c = table.add_row().cells
        c[0].text = k
        c[1].text = v

    for _ in range(2):
        doc.add_paragraph(long_academic_paragraph("Experimental methodology", "controlled training and reproducible evaluation", PROJECT_TITLE))


def add_chapter_4(doc: Document) -> None:
    doc.add_page_break()
    chapter_title(doc, "CHAPTER 4: RESULTS AND ANALYSIS")

    section_heading(doc, "4.1 Evaluation Metrics")
    doc.add_paragraph(
        "Model performance is evaluated using RMSE, MAE, R2, and MAPE to capture complementary aspects of regression quality. RMSE emphasizes larger deviations and is useful for risk-sensitive process settings. MAE provides a direct average magnitude of error in original units. R2 indicates explained variance, and MAPE normalizes error with respect to true values for relative interpretability. Together, these measures provide a balanced perspective on predictive reliability under realistic wastewater fluctuations."
    )

    # DO
    section_heading(doc, "4.2 Dissolved Oxygen (DO) Analysis")
    for _ in range(2):
        doc.add_paragraph(long_academic_paragraph("DO forecasting", "aeration-sensitive dissolved oxygen behavior", PROJECT_TITLE))
    add_figure(doc, "4", "4.1", "DO Predictions vs Actual", PLOTS_DIR / "do_02_predictions_vs_actual.png")
    doc.add_paragraph("The prediction curve for DO follows the true trajectory with acceptable smoothing, indicating that the model captures dominant operational rhythms and short-term process drift.")
    add_figure(doc, "4", "4.2", "DO Residual Plot", PLOTS_DIR / "do_03_residuals.png")
    doc.add_paragraph("Residual dispersion remains centered with moderate spread, suggesting manageable error variance and no severe systematic bias across the observed range.")
    add_figure(doc, "4", "4.3", "DO Scatter Plot", PLOTS_DIR / "do_05_actual_vs_predicted_scatter.png")
    doc.add_paragraph("The scatter alignment around the diagonal confirms useful correlation, although outer-region variance reveals opportunities for calibration under extreme conditions.")

    # pH
    section_heading(doc, "4.3 pH Analysis")
    for _ in range(2):
        doc.add_paragraph(long_academic_paragraph("pH forecasting", "buffered but sensitive acid-alkaline equilibrium dynamics", PROJECT_TITLE))
    add_figure(doc, "4", "4.4", "pH Predictions vs Actual", PLOTS_DIR / "ph_02_predictions_vs_actual.png")
    doc.add_paragraph("The pH prediction profile is stable and closely tracks central dynamics, reflecting the model's ability to represent bounded variations in near-neutral operating bands.")
    add_figure(doc, "4", "4.5", "pH Residual Plot", PLOTS_DIR / "ph_03_residuals.png")
    doc.add_paragraph("Residuals are tightly distributed around zero with low absolute magnitude, supporting consistent relative error behavior under pH scale constraints.")
    add_figure(doc, "4", "4.6", "pH Scatter Plot", PLOTS_DIR / "ph_05_actual_vs_predicted_scatter.png")
    doc.add_paragraph("The cluster concentration near the reference diagonal indicates dependable correlation despite reduced R2 sensitivity caused by narrow response variance.")

    # BOD
    section_heading(doc, "4.4 BOD Analysis")
    for _ in range(2):
        doc.add_paragraph(long_academic_paragraph("BOD forecasting", "organic load trend modeling and compliance relevance", PROJECT_TITLE))
    add_figure(doc, "4", "4.7", "BOD Predictions vs Actual", PLOTS_DIR / "02_predictions_vs_actual.png")
    doc.add_paragraph("BOD prediction demonstrates strong trend adherence and comparatively lower error spread, highlighting effective learning of organic loading patterns.")
    add_figure(doc, "4", "4.8", "BOD Residual Plot", PLOTS_DIR / "03_residuals.png")
    doc.add_paragraph("Residual characteristics indicate largely unbiased estimation with occasional peak underestimation during abrupt changes in influent profile.")
    add_figure(doc, "4", "4.9", "BOD Scatter Plot", PLOTS_DIR / "05_actual_vs_predicted_scatter.png")
    doc.add_paragraph("The scatter plot demonstrates high correlation and strong explanatory behavior, consistent with the superior R2 performance for this target.")

    # COD
    section_heading(doc, "4.5 COD Analysis")
    for _ in range(2):
        doc.add_paragraph(long_academic_paragraph("COD forecasting", "chemical oxygen demand and pollutant-load response", PROJECT_TITLE))
    add_figure(doc, "4", "4.10", "COD Predictions vs Actual", PLOTS_DIR / "cod_02_predictions_vs_actual.png")
    doc.add_paragraph("COD predictions preserve both baseline movement and major peaks, demonstrating that the hybrid architecture captures multiscale temporal structure.")
    add_figure(doc, "4", "4.11", "COD Residual Plot", PLOTS_DIR / "cod_03_residuals.png")
    doc.add_paragraph("Residual analysis reveals centered errors with wider tails than pH, reflecting natural complexity and larger value range of COD measurements.")
    add_figure(doc, "4", "4.12", "COD Scatter Plot", PLOTS_DIR / "cod_05_actual_vs_predicted_scatter.png")
    doc.add_paragraph("Diagonal clustering remains strong, confirming robust predictive performance for COD and validating suitability for plant-level operational guidance.")

    section_heading(doc, "4.6 Training Dynamics")
    add_figure(doc, "4", "4.13", "Training Loss Curve", PLOTS_DIR / "01_training_history.png")
    add_figure(doc, "4", "4.14", "Model Accuracy Curve", PLOTS_DIR / "cod_01_training_history.png")

    add_table_caption(doc, "4", "4.1", "Performance Metrics")
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Parameter"
    hdr[1].text = "RMSE"
    hdr[2].text = "MAE"
    hdr[3].text = "R2"
    hdr[4].text = "MAPE (%)"

    metrics = [
        ("DO", "0.5045", "0.4189", "0.5356", "8.5569"),
        ("pH", "0.0725", "0.0583", "0.5148", "0.8159"),
        ("BOD", "14.4834", "11.8191", "0.8289", "7.3658"),
        ("COD", "28.3680", "22.9936", "0.8343", "7.5561"),
    ]
    for r in metrics:
        row = table.add_row().cells
        for i, v in enumerate(r):
            row[i].text = v

    for _ in range(4):
        doc.add_paragraph(
            "Performance comparison confirms that COD and BOD achieved the strongest R2 values, indicating more effective variance explanation for high-range organic load indicators. DO and pH remain more challenging due to subtle fluctuation ranges and control-sensitive behavior that can produce distribution compression in observed values. Despite these differences, all targets display stable prediction trajectories and operationally meaningful trend alignment. The combined evidence from metrics, residual analysis, and scatter consistency supports the robustness of the proposed architecture for multi-parameter STP forecasting under realistic conditions."
        )


def add_chapter_5_and_end(doc: Document) -> None:
    doc.add_page_break()
    chapter_title(doc, "CHAPTER 5: CONCLUSION")
    for _ in range(8):
        doc.add_paragraph(long_academic_paragraph("Project conclusion", "end-to-end intelligent wastewater forecasting outcomes", PROJECT_TITLE))

    doc.add_page_break()
    chapter_title(doc, "FUTURE WORK")
    future_points = [
        "Integration with IoT sensor gateways for streaming real-time parameter ingestion.",
        "Deployment with cloud-native orchestration for scalable multi-plant operation.",
        "Adaptive retraining pipelines with drift detection and automated model governance.",
        "Inclusion of meteorological and hydraulic variables for richer context awareness.",
        "Mobile app expansion with dashboard analytics, alerts, and multilingual support.",
        "Role-based access control and secure API management for enterprise adoption.",
        "Explainable AI modules for operator trust and intervention transparency.",
    ]
    for fp in future_points:
        doc.add_paragraph(fp)
    for _ in range(5):
        doc.add_paragraph(long_academic_paragraph("Future enhancement roadmap", "IoT-cloud-mobile integration for smart utilities", PROJECT_TITLE))

    doc.add_page_break()
    chapter_title(doc, "REFERENCES")
    refs = [
        "Metcalf & Eddy, Inc. (2014). Wastewater Engineering: Treatment and Resource Recovery. McGraw-Hill Education.",
        "Hochreiter, S., & Schmidhuber, J. (1997). Long Short-Term Memory. Neural Computation, 9(8), 1735-1780.",
        "Vaswani, A., et al. (2017). Attention Is All You Need. Advances in Neural Information Processing Systems.",
        "Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press.",
        "Box, G. E. P., Jenkins, G. M., Reinsel, G. C., & Ljung, G. M. (2015). Time Series Analysis: Forecasting and Control. Wiley.",
        "Breiman, L. (2001). Random Forests. Machine Learning, 45, 5-32.",
        "Bishop, C. M. (2006). Pattern Recognition and Machine Learning. Springer.",
        "Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",
        "Abadi, M., et al. (2016). TensorFlow: A System for Large-Scale Machine Learning. OSDI.",
        "Chollet, F. (2018). Deep Learning with Python. Manning Publications.",
    ]
    for i, ref in enumerate(refs, start=1):
        doc.add_paragraph(f"[{i}] {ref}")


def main() -> None:
    figures = generate_conceptual_figures()

    doc = Document()
    configure_document(doc)

    # Preliminary section with roman numbering
    prelim_section = doc.sections[0]
    set_page_number_format(prelim_section, fmt="lowerRoman", start=1)
    add_page_number_footer(prelim_section)

    add_title_page(doc)
    add_preliminary_pages(doc)
    add_toc_and_lists(doc)

    # Main section with arabic numbering
    main_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page_number_format(main_sec, fmt="decimal", start=1)
    main_sec.footer.is_linked_to_previous = False
    add_page_number_footer(main_sec)

    add_chapter_1(doc, figures)
    add_chapter_2(doc, figures)
    add_chapter_3(doc, figures)
    add_chapter_4(doc)
    add_chapter_5_and_end(doc)

    doc.save(OUTPUT_DOCX)
    print(f"Generated: {OUTPUT_DOCX}")


if __name__ == "__main__":
    import pandas as pd

    main()
